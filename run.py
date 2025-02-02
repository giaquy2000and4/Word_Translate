import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from googletrans import Translator
from docx import Document

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    if file_path:
        entry_file_path.delete(0, tk.END)
        entry_file_path.insert(0, file_path)

def translate_and_save():
    file_path = entry_file_path.get()
    if not file_path:
        messagebox.showerror("Error", "Please select a Word file.")
        return

    try:
        document = Document(file_path)
        text = " ".join([para.text for para in document.paragraphs if para.text])

        translator = Translator()
        translated_text = ""
        chunk_size = 500
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i + chunk_size]
            translated_chunk = translator.translate(chunk, dest=combo_language.get()).text
            translated_text += translated_chunk + "\n"

        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
        if save_path:
            new_doc = Document()
            new_doc.add_paragraph(translated_text)
            new_doc.save(save_path)
            messagebox.showinfo("Success", "Translated file saved successfully.")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# Create root window
root = tk.Tk()
root.title("Word File Translator")
root.configure(bg="#2E2E2E")

# Frame for inputs
frame = tk.Frame(root, padx=10, pady=10, bg="#2E2E2E")
frame.pack(padx=10, pady=10)

# File selection
label_file = tk.Label(frame, text="Select Word File:", bg="#2E2E2E", fg="white")
label_file.grid(row=0, column=0, sticky="e")

entry_file_path = tk.Entry(frame, width=50, bg="#4E4E4E", fg="white", insertbackground="white")
entry_file_path.grid(row=0, column=1, padx=5)

btn_browse = tk.Button(frame, text="Browse", command=select_file, bg="#5A5A5A", fg="white")
btn_browse.grid(row=0, column=2, padx=5)

# Language selection
label_language = tk.Label(frame, text="Target Language:", bg="#2E2E2E", fg="white")
label_language.grid(row=1, column=0, sticky="e")

languages = {
    "Vietnamese": "vi",
    "English": "en",
    "French": "fr",
    "Spanish": "es",
    "German": "de",
    "Chinese": "zh-cn",
    "Japanese": "ja",
    "Korean": "ko"
}

combo_language = ttk.Combobox(frame, values=list(languages.keys()), state="readonly")
combo_language.grid(row=1, column=1, sticky="w")
combo_language.set("Vietnamese")

# Translate button
btn_translate = tk.Button(root, text="Translate and Save", command=translate_and_save, padx=10, pady=5, bg="#5A5A5A", fg="white")
btn_translate.pack(pady=10)

root.mainloop()
